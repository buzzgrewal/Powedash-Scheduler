import base64
import io
import json
import os
import re
import uuid
import smtplib
from email.message import EmailMessage
from datetime import datetime, timedelta, timezone, date, time
from typing import List, Dict, Any, Optional, Tuple

import fitz  # PyMuPDF
from PIL import Image
import streamlit as st

# --- Optional OpenAI (kept for PDF parsing flow) ---
try:
    from openai import OpenAI
except Exception:
    OpenAI = None  # type: ignore

from graph_client import GraphClient, GraphConfig, GraphAPIError, GraphAuthError
from audit_log import AuditLog, LogLevel, log_structured
from ics_utils import ICSInvite, stable_uid, ICSValidationError
from timezone_utils import to_utc, from_utc, iso_utc, is_valid_timezone, safe_zoneinfo


# ----------------------------
# Input Validation
# ----------------------------
import re as _re
from typing import Tuple as _Tuple

# Email regex (RFC 5322 simplified)
_EMAIL_REGEX = _re.compile(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$')

# Date/time patterns from OpenAI output
_DATE_REGEX = _re.compile(r'^\d{4}-\d{2}-\d{2}$')
_TIME_REGEX = _re.compile(r'^\d{2}:\d{2}$')


class ValidationError(ValueError):
    """Raised when input validation fails."""
    def __init__(self, field: str, message: str):
        self.field = field
        self.message = message
        super().__init__(f"{field}: {message}")


def validate_email(email: str, field_name: str = "email") -> str:
    """Validate email format. Returns cleaned email or raises ValidationError."""
    if not email:
        raise ValidationError(field_name, "Email is required")
    email = email.strip().lower()
    if not _EMAIL_REGEX.match(email):
        raise ValidationError(field_name, f"Invalid email format: {email}")
    if len(email) > 254:  # RFC 5321 limit
        raise ValidationError(field_name, "Email too long (max 254 characters)")
    return email


def validate_email_optional(email: Optional[str], field_name: str = "email") -> Optional[str]:
    """Validate email if provided, return None if empty."""
    if not email or not email.strip():
        return None
    return validate_email(email, field_name)


def validate_slot(slot: dict) -> _Tuple[str, str, str]:
    """Validate slot dict from OpenAI parsing. Returns (date, start, end) tuple."""
    if not isinstance(slot, dict):
        raise ValidationError("slot", "Slot must be a dictionary")

    date = slot.get("date", "")
    start = slot.get("start", "")
    end = slot.get("end", "")

    if not _DATE_REGEX.match(date):
        raise ValidationError("slot.date", f"Invalid date format: {date}. Expected YYYY-MM-DD")
    if not _TIME_REGEX.match(start):
        raise ValidationError("slot.start", f"Invalid start time format: {start}. Expected HH:MM")
    if end and not _TIME_REGEX.match(end):
        raise ValidationError("slot.end", f"Invalid end time format: {end}. Expected HH:MM")

    return date, start, end


# ----------------------------
# Configuration helpers
# ----------------------------
def get_secret(key: str, default: Any = None) -> Any:
    # st.secrets behaves like a dict on Streamlit Cloud; local dev can use env vars too.
    try:
        if key in st.secrets:
            return st.secrets.get(key, default)
    except Exception:
        pass
    return os.getenv(key.upper(), default)


def get_default_timezone() -> str:
    return get_secret("default_timezone", "Europe/London")


def get_audit_log_path() -> str:
    return get_secret("audit_log_path", "audit_log.db")


def get_graph_config() -> Optional[GraphConfig]:
    tenant_id = get_secret("graph_tenant_id")
    client_id = get_secret("graph_client_id")
    client_secret = get_secret("graph_client_secret")
    scheduler_mailbox = get_secret("graph_scheduler_mailbox", "scheduling@powerdashhr.com")

    if not (tenant_id and client_id and client_secret and scheduler_mailbox):
        return None
    return GraphConfig(
        tenant_id=str(tenant_id),
        client_id=str(client_id),
        client_secret=str(client_secret),
        scheduler_mailbox=str(scheduler_mailbox),
    )


def graph_enabled() -> bool:
    return get_graph_config() is not None


def get_openai_client() -> Optional[Any]:
    """
    Create an OpenAI client using Streamlit secrets or environment variable.
    Keeps backward compatibility with older secret key names.
    """
    api_key = get_secret("openai_api_key") or get_secret("OPENAI_API_KEY")
    if not api_key:
        st.warning("OpenAI API key not found. PDF availability parsing will be limited.")
        return None
    if OpenAI is None:
        st.warning("OpenAI SDK not available (openai). PDF availability parsing will be limited.")
        return None
    return OpenAI(api_key=api_key)


# ----------------------------
# PDF / image parsing helpers (existing behavior)
# ----------------------------
def image_to_base64(image: Image.Image) -> str:
    buffered = io.BytesIO()
    image.save(buffered, format="PNG")
    return base64.b64encode(buffered.getvalue()).decode("utf-8")


def parse_slots_from_image(image: Image.Image) -> List[Dict[str, str]]:
    """
    Use OpenAI vision to parse free/busy calendar images into slots.
    Expected JSON format:
    [
      {"date": "2025-12-03", "start": "09:00", "end": "09:30"},
      ...
    ]
    """
    client = get_openai_client()
    if not client:
        return []

    prompt = (
        "You are extracting FREE time slots from a calendar screenshot.\n"
        "Return ONLY valid JSON (no markdown) as a list of objects with keys:\n"
        "  - date (YYYY-MM-DD)\n"
        "  - start (HH:MM in 24-hour format)\n"
        "  - end (HH:MM in 24-hour format)\n"
        "  - inferred_tz (timezone abbreviation if visible, e.g. 'PST', 'EST', 'GMT', or null if not visible)\n\n"
        "Look for timezone indicators in:\n"
        "  - Calendar headers or footers\n"
        "  - Corner labels (e.g., 'Times shown in PST')\n"
        "  - Time displays with timezone suffix (e.g., '2:00 PM EST')\n"
        "  - UTC offset indicators (e.g., 'GMT-8')\n\n"
        "Only include free slots that are at least 30 minutes.\n"
        "If no slots found, return an empty list []."
    )

    b64 = image_to_base64(image)

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0,
            messages=[
                {"role": "system", "content": "You are a helpful assistant that returns strict JSON."},
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
                    ],
                },
            ],
        )
        content = resp.choices[0].message.content.strip() if resp.choices else ""
        # Strip code fences if present
        if content.startswith("```"):
            content = content.strip("`")
            if "\n" in content:
                content = content.split("\n", 1)[1].strip()

        slots = json.loads(content) if content else []
        valid_slots = []
        for s in slots:
            if isinstance(s, dict) and all(k in s for k in ("date", "start", "end")):
                slot_data = {
                    "date": str(s["date"]),
                    "start": str(s["start"]),
                    "end": str(s["end"]),
                }
                # Preserve inferred timezone if present
                if s.get("inferred_tz"):
                    slot_data["inferred_tz"] = str(s["inferred_tz"])
                valid_slots.append(slot_data)
        return valid_slots
    except json.JSONDecodeError as e:
        st.error(f"OpenAI returned invalid JSON: {e}")
        log_structured(
            LogLevel.ERROR,
            f"OpenAI JSON parse error: {e}",
            action="parse_slots_openai",
            error_type="json_decode_error",
        )
        return []
    except Exception as e:
        st.error(f"Failed to parse availability via OpenAI: {e}")
        log_structured(
            LogLevel.ERROR,
            f"OpenAI vision API error: {e}",
            action="parse_slots_openai",
            error_type=type(e).__name__,
            exc_info=True,
        )
        return []


def pdf_to_images(pdf_bytes: bytes, max_pages: int = 3) -> List[Image.Image]:
    """Convert PDF to images. Returns empty list on error instead of crashing."""
    images: List[Image.Image] = []
    doc = None
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for i in range(min(len(doc), max_pages)):
            try:
                page = doc.load_page(i)
                pix = page.get_pixmap(dpi=200)
                img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
                images.append(img)
            except Exception as e:
                log_structured(
                    LogLevel.WARNING,
                    f"Failed to process PDF page {i}: {e}",
                    action="pdf_page_process",
                    error_type="pdf_error",
                )
    except Exception as e:
        st.error(f"Failed to open PDF: {e}")
        log_structured(
            LogLevel.ERROR,
            f"Failed to open PDF: {e}",
            action="pdf_open",
            error_type="pdf_error",
            exc_info=True,
        )
    finally:
        if doc:
            doc.close()
    return images


def docx_to_text(docx_bytes: bytes) -> str:
    """
    Extract text from a Word document including paragraphs and tables.
    Returns empty string on error instead of crashing.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        st.warning("python-docx not installed. Word document parsing unavailable.")
        log_structured(
            LogLevel.ERROR,
            "python-docx not installed",
            action="docx_import",
            error_type="import_error",
        )
        return ""

    try:
        doc = DocxDocument(io.BytesIO(docx_bytes))
        text_parts: List[str] = []

        # Extract paragraphs
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                text_parts.append(para_text)

        # Extract tables (important for calendar/availability data)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)
    except Exception as e:
        st.error(f"Failed to read Word document: {e}")
        log_structured(
            LogLevel.ERROR,
            f"Failed to read Word document: {e}",
            action="docx_read",
            error_type=type(e).__name__,
            exc_info=True,
        )
        return ""


def docx_extract_images(docx_bytes: bytes, max_images: int = 5) -> List[Image.Image]:
    """
    Extract embedded images from a Word document.
    Returns empty list on error instead of crashing.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return []

    images: List[Image.Image] = []
    try:
        doc = DocxDocument(io.BytesIO(docx_bytes))

        # Access the document's related parts to find images
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    img = Image.open(io.BytesIO(image_data)).convert("RGB")
                    images.append(img)
                    if len(images) >= max_images:
                        break
                except Exception as e:
                    log_structured(
                        LogLevel.WARNING,
                        f"Failed to extract image from docx: {e}",
                        action="docx_image_extract",
                        error_type="image_error",
                    )
                    continue

        return images
    except Exception as e:
        log_structured(
            LogLevel.WARNING,
            f"Failed to extract images from Word document: {e}",
            action="docx_image_extract",
            error_type=type(e).__name__,
        )
        return []


def parse_slots_from_text(text: str) -> List[Dict[str, str]]:
    """
    Use OpenAI to parse free/busy text into slots.
    Expected JSON format:
    [
      {"date": "2025-12-03", "start": "09:00", "end": "09:30"},
      ...
    ]
    """
    if not text or not text.strip():
        return []

    client = get_openai_client()
    if not client:
        return []

    # Get current year for inference
    current_year = datetime.now().year

    prompt = f"""You are extracting FREE/AVAILABLE time slots from text describing someone's availability.

IMPORTANT RULES:
1. Only extract slots explicitly marked as FREE, AVAILABLE, or OPEN
2. Do NOT include busy/blocked/unavailable times
3. Convert all dates to YYYY-MM-DD format
4. Convert all times to 24-hour HH:MM format
5. If year is not specified, assume {current_year}
6. If end time is not specified, assume 1 hour duration
7. Only include slots that are at least 30 minutes

DATE FORMAT EXAMPLES:
- "Monday Dec 3" -> "{current_year}-12-03"
- "12/03/2025" -> "2025-12-03"
- "3rd December" -> "{current_year}-12-03"
- "Dec 3, 2025" -> "2025-12-03"

TIME FORMAT EXAMPLES:
- "9am-10am" -> start: "09:00", end: "10:00"
- "09:00-10:00" -> start: "09:00", end: "10:00"
- "9:00 AM to 10:00 AM" -> start: "09:00", end: "10:00"
- "2pm-3:30pm" -> start: "14:00", end: "15:30"

Return ONLY valid JSON as a list of objects with keys: date, start, end.
If no free slots found, return an empty list [].

TEXT TO PARSE:
{text}"""

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0,
            messages=[
                {"role": "system", "content": "You are a helpful assistant that returns strict JSON. Never include markdown formatting."},
                {"role": "user", "content": prompt},
            ],
        )
        content = resp.choices[0].message.content.strip() if resp.choices else ""

        # Strip code fences if present (same pattern as parse_slots_from_image)
        if content.startswith("```"):
            content = content.strip("`")
            if "\n" in content:
                content = content.split("\n", 1)[1].strip()

        slots = json.loads(content) if content else []
        valid_slots = []
        for s in slots:
            if isinstance(s, dict) and all(k in s for k in ("date", "start", "end")):
                valid_slots.append({
                    "date": str(s["date"]),
                    "start": str(s["start"]),
                    "end": str(s["end"])
                })
        return valid_slots
    except json.JSONDecodeError as e:
        st.error(f"OpenAI returned invalid JSON: {e}")
        log_structured(
            LogLevel.ERROR,
            f"OpenAI JSON parse error: {e}",
            action="parse_slots_text_openai",
            error_type="json_decode_error",
        )
        return []
    except Exception as e:
        st.error(f"Failed to parse availability via OpenAI: {e}")
        log_structured(
            LogLevel.ERROR,
            f"OpenAI text API error: {e}",
            action="parse_slots_text_openai",
            error_type=type(e).__name__,
            exc_info=True,
        )
        return []


def ensure_session_state() -> None:
    defaults = {
        "slots": [],
        "last_graph_event_id": "",
        "last_teams_join_url": "",
        "last_invite_uid": "",
        "last_invite_ics_bytes": b"",
        "selected_timezone": get_default_timezone(),
        "candidate_timezone": get_default_timezone(),
        "duration_minutes": 30,
        # Panel interview support
        "panel_interviewers": [],  # List of {id, name, email, file, slots, timezone}
        "next_interviewer_id": 1,  # Auto-increment for unique widget keys
        "slot_filter_mode": "all_available",  # "all_available" | "any_n" | "show_all"
        "slot_filter_min_n": 1,  # Minimum N for "any_n" mode
        "computed_intersections": [],  # Intersection slots with availability metadata
        "editing_slot_index": None,  # Track which slot is being edited: (interviewer_idx, slot_idx) or None
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def format_slot_label(slot: Dict[str, str]) -> str:
    return f"{slot['date']} {slot['start']}–{slot['end']}"


def _merge_slots(manual_slots: List[Dict], uploaded_slots: List[Dict]) -> List[Dict]:
    """Merge slots, preferring manual over uploaded for duplicates."""
    seen = {}
    for s in manual_slots:
        key = (s["date"], s["start"], s["end"])
        seen[key] = s
    for s in uploaded_slots:
        key = (s["date"], s["start"], s["end"])
        if key not in seen:
            seen[key] = s
    return list(seen.values())


def _add_manual_slot(interviewer_idx: int, slot_date, start_time, end_time) -> bool:
    """Add a manually entered slot with validation. Returns True if successful."""
    from datetime import date as date_type, time as time_type

    errors = []

    # Validate end time is after start time
    if end_time <= start_time:
        errors.append("End time must be after start time")

    # Validate minimum duration (30 minutes)
    start_dt = datetime.combine(date.today(), start_time)
    end_dt = datetime.combine(date.today(), end_time)
    duration_minutes = (end_dt - start_dt).seconds // 60
    if duration_minutes < 30:
        errors.append("Slot must be at least 30 minutes")

    # Validate not in the past
    if slot_date < date.today():
        errors.append("Cannot add slots in the past")

    if errors:
        for err in errors:
            st.error(err)
        return False

    # Create slot in standard format
    new_slot = {
        "date": slot_date.strftime("%Y-%m-%d"),
        "start": start_time.strftime("%H:%M"),
        "end": end_time.strftime("%H:%M"),
        "source": "manual",
    }

    # Get existing slots for this interviewer
    interviewers = st.session_state.get("panel_interviewers", [])
    if interviewer_idx >= len(interviewers):
        st.error("Invalid interviewer index")
        return False

    existing_slots = interviewers[interviewer_idx].get("slots", [])

    # Check for duplicates
    slot_key = (new_slot["date"], new_slot["start"], new_slot["end"])
    for s in existing_slots:
        if (s["date"], s["start"], s["end"]) == slot_key:
            st.warning("This slot already exists")
            return False

    existing_slots.append(new_slot)
    st.session_state["panel_interviewers"][interviewer_idx]["slots"] = existing_slots
    st.success(f"Added slot: {format_slot_label(new_slot)}")
    return True


def _delete_interviewer_slot(interviewer_idx: int, slot_idx: int) -> None:
    """Delete a slot by index from an interviewer's slots."""
    interviewers = st.session_state.get("panel_interviewers", [])
    if interviewer_idx >= len(interviewers):
        return

    slots = interviewers[interviewer_idx].get("slots", [])
    if 0 <= slot_idx < len(slots):
        deleted = slots.pop(slot_idx)
        st.session_state["panel_interviewers"][interviewer_idx]["slots"] = slots
        st.toast(f"Deleted: {format_slot_label(deleted)}")
        st.rerun()


def _render_interviewer_slots(interviewer_idx: int, interviewer_id: int) -> None:
    """Render editable list of current slots for an interviewer."""
    interviewers = st.session_state.get("panel_interviewers", [])
    if interviewer_idx >= len(interviewers):
        return

    slots = interviewers[interviewer_idx].get("slots", [])

    if not slots:
        st.info("No slots added yet. Use the form above or upload a calendar.")
        return

    st.markdown(f"**Current Slots ({len(slots)}):**")

    for idx, slot in enumerate(slots):
        col_label, col_edit, col_delete = st.columns([4, 1, 1])

        with col_label:
            source_badge = " manual" if slot.get("source") == "manual" else " uploaded"
            st.text(f"{source_badge} {format_slot_label(slot)}")

        with col_edit:
            if st.button("Edit", key=f"edit_slot_{interviewer_id}_{idx}"):
                st.session_state["editing_slot_index"] = (interviewer_idx, idx)
                st.rerun()

        with col_delete:
            if st.button("Del", key=f"del_slot_{interviewer_id}_{idx}"):
                _delete_interviewer_slot(interviewer_idx, idx)

    # Clear all button
    if len(slots) > 1:
        if st.button("Clear All Slots", key=f"clear_all_{interviewer_id}", type="secondary"):
            st.session_state["panel_interviewers"][interviewer_idx]["slots"] = []
            st.rerun()


def _render_manual_slot_form(interviewer_idx: int, interviewer_id: int) -> None:
    """Render the form to add a new manual slot."""
    st.caption("Add availability slots manually instead of uploading a calendar")

    col_date, col_start, col_end, col_btn = st.columns([2, 1.5, 1.5, 1])

    with col_date:
        slot_date = st.date_input(
            "Date",
            value=date.today(),
            key=f"manual_slot_date_{interviewer_id}",
            min_value=date.today(),
        )
    with col_start:
        slot_start = st.time_input(
            "Start",
            value=time(9, 0),
            key=f"manual_slot_start_{interviewer_id}",
        )
    with col_end:
        slot_end = st.time_input(
            "End",
            value=time(10, 0),
            key=f"manual_slot_end_{interviewer_id}",
        )
    with col_btn:
        st.write("")  # Vertical spacing
        if st.button("+ Add", key=f"add_manual_slot_{interviewer_id}", type="primary"):
            if _add_manual_slot(interviewer_idx, slot_date, slot_start, slot_end):
                st.rerun()


def _render_edit_slot_form(interviewer_idx: int, interviewer_id: int) -> None:
    """Render edit form when a slot is being edited."""
    edit_info = st.session_state.get("editing_slot_index")
    if edit_info is None:
        return

    edit_interviewer_idx, edit_slot_idx = edit_info

    # Only render if this is the interviewer being edited
    if edit_interviewer_idx != interviewer_idx:
        return

    interviewers = st.session_state.get("panel_interviewers", [])
    if interviewer_idx >= len(interviewers):
        st.session_state["editing_slot_index"] = None
        return

    slots = interviewers[interviewer_idx].get("slots", [])
    if edit_slot_idx >= len(slots):
        st.session_state["editing_slot_index"] = None
        return

    slot = slots[edit_slot_idx]

    st.markdown("---")
    st.markdown(f"**Editing:** {format_slot_label(slot)}")

    col_date, col_start, col_end = st.columns(3)
    with col_date:
        new_date = st.date_input(
            "Date",
            value=datetime.strptime(slot["date"], "%Y-%m-%d").date(),
            key=f"edit_slot_date_{interviewer_id}",
        )
    with col_start:
        new_start = st.time_input(
            "Start",
            value=datetime.strptime(slot["start"], "%H:%M").time(),
            key=f"edit_slot_start_{interviewer_id}",
        )
    with col_end:
        new_end = st.time_input(
            "End",
            value=datetime.strptime(slot["end"], "%H:%M").time(),
            key=f"edit_slot_end_{interviewer_id}",
        )

    col_save, col_cancel = st.columns(2)
    with col_save:
        if st.button("Save Changes", type="primary", key=f"save_edit_{interviewer_id}"):
            # Validate
            if new_end <= new_start:
                st.error("End time must be after start time")
            elif new_date < date.today():
                st.error("Cannot set date in the past")
            else:
                duration = (datetime.combine(date.today(), new_end) - datetime.combine(date.today(), new_start)).seconds // 60
                if duration < 30:
                    st.error("Slot must be at least 30 minutes")
                else:
                    # Update the slot
                    slots[edit_slot_idx] = {
                        "date": new_date.strftime("%Y-%m-%d"),
                        "start": new_start.strftime("%H:%M"),
                        "end": new_end.strftime("%H:%M"),
                        "source": slot.get("source", "manual"),
                    }
                    st.session_state["panel_interviewers"][interviewer_idx]["slots"] = slots
                    st.session_state["editing_slot_index"] = None
                    st.success("Slot updated!")
                    st.rerun()

    with col_cancel:
        if st.button("Cancel", key=f"cancel_edit_{interviewer_id}"):
            st.session_state["editing_slot_index"] = None
            st.rerun()

    st.markdown("---")


def extract_common_timezone(slots: List[Dict[str, str]]) -> Optional[str]:
    """
    Extract the most common inferred timezone from parsed slots.

    Returns IANA timezone or None if no timezone was inferred.
    """
    from collections import Counter
    from timezone_utils import infer_timezone_from_abbreviation

    tz_abbrevs = [s.get("inferred_tz") for s in slots if s.get("inferred_tz")]
    if not tz_abbrevs:
        return None

    # Get most common abbreviation
    most_common = Counter(tz_abbrevs).most_common(1)[0][0]

    # Convert to IANA timezone name
    iana_tz, matched, _ = infer_timezone_from_abbreviation(most_common)
    return iana_tz if matched else None


# ----------------------------
# Email helpers (existing, with updated secret key names)
# ----------------------------
def build_scheduling_email(role_title: str, recruiter_name: str, slots: List[Dict[str, str]]) -> str:
    slot_lines = "\n".join([f"- {format_slot_label(s)}" for s in slots]) if slots else "- (No slots extracted)"
    return f"""Hi there,

Thanks for your time. Please choose one of the following interview times for the role: {role_title}

Available slots:
{slot_lines}

Reply with your preferred option and we will confirm the invite.

Best regards,
{recruiter_name}
Talent Acquisition
"""


def _smtp_cfg() -> Optional[Dict[str, Any]]:
    # New keys (preferred)
    host = get_secret("smtp_host")
    port = get_secret("smtp_port")
    username = get_secret("smtp_username")
    password = get_secret("smtp_password")
    smtp_from = get_secret("smtp_from")

    # Backward-compat keys (older app)
    if not host:
        host = get_secret("smtp_server")

    if not (host and username and password):
        return None

    return {
        "host": str(host),
        "port": int(port or 587),
        "username": str(username),
        "password": str(password),
        "from": str(smtp_from or username),
    }


def send_email_smtp(
    subject: str,
    body: str,
    to_emails: List[str],
    cc_emails: Optional[List[str]] = None,
    attachment: Optional[Dict[str, Any]] = None,
) -> bool:
    cfg = _smtp_cfg()
    if not cfg:
        st.warning("SMTP is not configured in secrets; email send is disabled.")
        return False

    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = cfg["from"]
    msg["To"] = ", ".join([e for e in to_emails if e])
    if cc_emails:
        msg["Cc"] = ", ".join([e for e in cc_emails if e])
    msg.set_content(body)

    if attachment:
        msg.add_attachment(
            attachment["data"],
            maintype=attachment.get("maintype", "application"),
            subtype=attachment.get("subtype", "octet-stream"),
            filename=attachment.get("filename", "attachment.bin"),
        )

    try:
        with smtplib.SMTP(cfg["host"], cfg["port"]) as server:
            server.starttls()
            server.login(cfg["username"], cfg["password"])
            server.send_message(msg)
        return True
    except smtplib.SMTPAuthenticationError as e:
        st.error(f"SMTP authentication failed: {e}")
        log_structured(
            LogLevel.ERROR,
            f"SMTP authentication failed: {e}",
            action="smtp_send",
            error_type="smtp_auth_error",
        )
        return False
    except smtplib.SMTPException as e:
        st.error(f"SMTP send failed: {e}")
        log_structured(
            LogLevel.ERROR,
            f"SMTP send failed: {e}",
            action="smtp_send",
            error_type="smtp_error",
        )
        return False
    except Exception as e:
        st.error(f"SMTP send failed: {e}")
        log_structured(
            LogLevel.ERROR,
            f"SMTP send failed: {e}",
            action="smtp_send",
            error_type=type(e).__name__,
            exc_info=True,
        )
        return False


def send_email_graph(
    subject: str,
    body: str,
    to_emails: List[str],
    cc_emails: Optional[List[str]] = None,
    attachment: Optional[Dict[str, Any]] = None,
) -> bool:
    """Send email using Microsoft Graph API."""
    cfg = get_graph_config()
    if not cfg:
        st.warning("Graph is not configured. Add graph_tenant_id, graph_client_id, graph_client_secret, graph_scheduler_mailbox in Streamlit secrets.")
        return False

    try:
        client = GraphClient(cfg)
        graph_attachment = None
        if attachment:
            graph_attachment = {
                "name": attachment.get("filename", "attachment.bin"),
                "contentBytes": attachment.get("data"),
                "contentType": f"{attachment.get('maintype', 'application')}/{attachment.get('subtype', 'octet-stream')}",
            }
        client.send_mail(
            subject=subject,
            body=body,
            to_recipients=[e for e in to_emails if e],
            cc_recipients=[e for e in (cc_emails or []) if e] or None,
            content_type="Text",
            attachment=graph_attachment,
        )
        return True
    except Exception as e:
        st.error(f"Graph email send failed: {e}")
        return False


def fetch_unread_emails_graph() -> Tuple[List[Dict[str, Any]], Optional[str], bool]:
    """
    Fetch unread emails from scheduler mailbox via Microsoft Graph API.
    Returns (emails, error_message, is_configured) tuple.
    - error_message is None on success
    - is_configured is False if Graph credentials are missing

    Uses the same Graph credentials as calendar operations.
    """
    cfg = get_graph_config()
    if not cfg:
        return [], None, False  # Graph not configured

    try:
        from graph_client import GraphClient
        client = GraphClient(cfg)
        messages = client.fetch_unread_messages(top=50)

        emails: List[Dict[str, Any]] = []
        for msg in messages:
            from_addr = ""
            from_data = msg.get("from", {})
            if from_data:
                email_addr = from_data.get("emailAddress", {})
                from_addr = email_addr.get("address", "")

            # Get body content (prefer text, fall back to HTML)
            body_content = msg.get("bodyPreview", "")
            body_data = msg.get("body", {})
            if body_data and body_data.get("content"):
                body_content = body_data.get("content", "")
                # Strip HTML tags if content type is HTML
                if body_data.get("contentType") == "html":
                    import re
                    body_content = re.sub(r'<[^>]+>', '', body_content)
                    body_content = body_content.strip()

            emails.append({
                "id": msg.get("id", ""),
                "from": from_addr,
                "subject": msg.get("subject", ""),
                "date": msg.get("receivedDateTime", ""),
                "body": body_content,
            })

        return emails, None, True  # Success, configured

    except GraphAuthError as e:
        log_structured(
            LogLevel.ERROR,
            f"Graph authentication failed: {e}",
            action="graph_fetch_messages",
            error_type="graph_auth_error",
        )
        return [], f"Graph authentication failed: {e}", True
    except GraphAPIError as e:
        log_structured(
            LogLevel.ERROR,
            f"Graph API error: {e}",
            action="graph_fetch_messages",
            error_type="graph_api_error",
            details={"status_code": e.status_code},
        )
        return [], f"Graph API error: {e}", True
    except Exception as e:
        log_structured(
            LogLevel.ERROR,
            f"Failed to fetch emails via Graph: {e}",
            action="graph_fetch_messages",
            error_type="graph_error",
            exc_info=True,
        )
        return [], f"Failed to fetch emails: {e}", True


def detect_slot_choice_from_text(text: str, slots: List[Dict[str, str]]) -> Optional[Dict[str, str]]:
    """
    Heuristic: find a slot label or date+time mention in a reply.
    """
    t = (text or "").lower()
    for s in slots:
        label = format_slot_label(s).lower()
        if label in t:
            return s

    # fallback: look for YYYY-MM-DD and HH:MM
    m_date = re.search(r"\b(20\d{2}-\d{2}-\d{2})\b", t)
    m_time = re.search(r"\b(\d{1,2}:\d{2})\b", t)
    if m_date and m_time:
        date = m_date.group(1)
        start = m_time.group(1).zfill(5)
        for s in slots:
            if s["date"] == date and s["start"] == start:
                return s
        return {"date": date, "start": start, "end": ""}

    return None


# ----------------------------
# Graph + ICS helpers
# ----------------------------
def _make_graph_client() -> Optional[GraphClient]:
    cfg = get_graph_config()
    if not cfg:
        return None
    return GraphClient(cfg)


def _graph_event_payload(
    *,
    subject: str,
    body_html: str,
    start_local: datetime,
    end_local: datetime,
    time_zone: str,
    attendees: List[Tuple[str, str]],
    is_teams: bool,
    location: str,
) -> Dict[str, Any]:
    payload: Dict[str, Any] = {
        "subject": subject,
        "body": {"contentType": "HTML", "content": body_html},
        "start": {"dateTime": start_local.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": time_zone},
        "end": {"dateTime": end_local.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": time_zone},
        "attendees": [{"emailAddress": {"address": e, "name": n or e}, "type": "required"} for (e, n) in attendees],
    }

    if is_teams:
        payload["isOnlineMeeting"] = True
        payload["onlineMeetingProvider"] = "teamsForBusiness"
        payload["location"] = {"displayName": "Microsoft Teams"}
    else:
        payload["location"] = {"displayName": location or "Interview"}

    return payload


def _build_ics(
    *,
    organizer_email: str,
    organizer_name: str,
    attendee_emails: List[str],
    summary: str,
    description: str,
    dtstart_utc: datetime,
    dtend_utc: datetime,
    location: str,
    url: str,
    uid_hint: str,
    display_timezone: str = "UTC",
) -> bytes:
    uid = stable_uid(uid_hint, organizer_email, ",".join(attendee_emails), dtstart_utc.isoformat())
    inv = ICSInvite(
        uid=uid,
        dtstart_utc=dtstart_utc,
        dtend_utc=dtend_utc,
        summary=summary,
        description=description,
        organizer_email=organizer_email,
        organizer_name=organizer_name,
        attendee_emails=attendee_emails,
        location=location,
        url=url,
        display_timezone=display_timezone,
    )
    return inv.to_ics()


# ----------------------------
# Streamlit UI
# ----------------------------
def main() -> None:
    st.set_page_config(page_title="PowerDash Interview Scheduler", layout="wide")
    ensure_session_state()

    audit = AuditLog(get_audit_log_path())

    st.title("PowerDash Interview Scheduler")

    tab_new, tab_inbox, tab_invites, tab_audit, tab_diag = st.tabs(
        ["New Scheduling Request", "Scheduler Inbox", "Calendar Invites", "Audit Log", "Graph Diagnostics"]
    )

    # ========= TAB: New Scheduling Request =========
    with tab_new:
        st.subheader("New Scheduling Request")

        col_left, col_center, col_right = st.columns([1.2, 1.5, 1.2], gap="large")

        with col_left:
            st.markdown("#### Hiring Manager & Recruiter")
            role_title = st.text_input("Role Title", key="role_title")
            hiring_manager_name = st.text_input("Hiring Manager Name", key="hm_name")
            hiring_manager_email = st.text_input("Hiring Manager Email (required)", key="hm_email")
            recruiter_name = st.text_input("Recruiter Name", key="rec_name")
            recruiter_email = st.text_input("Recruiter Email (optional attendee)", key="rec_email")
            scheduler_mailbox = get_secret("graph_scheduler_mailbox", "scheduling@powerdashhr.com")
            st.text_input("Recruiter / Scheduling Mailbox Email", value=str(scheduler_mailbox), disabled=True)

        with col_center:
            st.markdown("#### Interviewer Availability")

            # Ensure at least one interviewer exists
            if not st.session_state.get("panel_interviewers"):
                new_id = st.session_state["next_interviewer_id"]
                st.session_state["next_interviewer_id"] = new_id + 1
                st.session_state["panel_interviewers"] = [{
                    "id": new_id,
                    "name": "",
                    "email": "",
                    "file": None,
                    "slots": [],
                    "timezone": st.session_state["selected_timezone"],
                }]

            interviewers = st.session_state["panel_interviewers"]

            # Render each interviewer
            for idx, interviewer in enumerate(interviewers):
                with st.container(border=True):
                    cols = st.columns([3, 3, 1])
                    with cols[0]:
                        name = st.text_input(
                            "Name",
                            value=interviewer.get("name", ""),
                            key=f"interviewer_name_{interviewer['id']}",
                            placeholder="e.g., John Smith"
                        )
                        interviewers[idx]["name"] = name
                    with cols[1]:
                        email = st.text_input(
                            "Email",
                            value=interviewer.get("email", ""),
                            key=f"interviewer_email_{interviewer['id']}",
                            placeholder="john@company.com"
                        )
                        interviewers[idx]["email"] = email
                    with cols[2]:
                        st.write("")  # Spacing
                        # Remove button (disabled if only 1 interviewer)
                        if len(interviewers) > 1:
                            if st.button("Remove", key=f"remove_{interviewer['id']}"):
                                st.session_state["panel_interviewers"] = [
                                    i for i in interviewers if i["id"] != interviewer["id"]
                                ]
                                st.rerun()

                    # File uploader
                    uploaded = st.file_uploader(
                        f"Calendar ({interviewer.get('name') or f'Interviewer {idx+1}'})",
                        type=["pdf", "png", "jpg", "jpeg", "docx"],
                        key=f"file_{interviewer['id']}",
                    )
                    interviewers[idx]["file"] = uploaded

                    # Show slot count with breakdown
                    slot_count = len(interviewer.get("slots", []))
                    manual_count = len([s for s in interviewer.get("slots", []) if s.get("source") == "manual"])
                    uploaded_count = slot_count - manual_count
                    if slot_count > 0:
                        if manual_count > 0 and uploaded_count > 0:
                            st.caption(f"{slot_count} slot(s) ({manual_count} manual, {uploaded_count} uploaded)")
                        elif manual_count > 0:
                            st.caption(f"{slot_count} manual slot(s)")
                        else:
                            st.caption(f"{slot_count} uploaded slot(s)")

                    # Manual slot entry expander
                    with st.expander("Manual Slot Entry", expanded=False):
                        _render_manual_slot_form(idx, interviewer["id"])
                        _render_edit_slot_form(idx, interviewer["id"])
                        _render_interviewer_slots(idx, interviewer["id"])

            st.session_state["panel_interviewers"] = interviewers

            # Add interviewer button
            if st.button("+ Add Interviewer", key="add_interviewer_btn"):
                new_id = st.session_state["next_interviewer_id"]
                st.session_state["next_interviewer_id"] = new_id + 1
                st.session_state["panel_interviewers"].append({
                    "id": new_id,
                    "name": "",
                    "email": "",
                    "file": None,
                    "slots": [],
                    "timezone": st.session_state["selected_timezone"],
                })
                st.rerun()

            st.markdown("---")

            st.session_state["duration_minutes"] = st.number_input(
                "Interview duration (minutes)", min_value=15, max_value=240, step=15, value=int(st.session_state["duration_minutes"])
            )
            tz_name = st.selectbox(
                "Display timezone",
                options=_common_timezones(),
                index=_common_timezones().index(st.session_state["selected_timezone"])
                if st.session_state["selected_timezone"] in _common_timezones()
                else 0,
                key="selected_timezone",
            )

            # Real-time clock showing current time in selected timezone vs system timezone
            from timezone_utils import from_utc
            now_utc = datetime.now(timezone.utc)
            now_system = datetime.now().astimezone()  # System local time
            system_tz_name = now_system.strftime("%Z")  # e.g., "PST", "GMT"

            try:
                now_selected = from_utc(now_utc, tz_name)
                selected_time = now_selected.strftime("%I:%M %p %Z")
                system_time = now_system.strftime("%I:%M %p %Z")

                st.caption(f"**{tz_name}**: {selected_time} | **Your system ({system_tz_name})**: {system_time}")
            except:
                pass

            parse_btn = st.button("Parse All Availability", type="primary")

            if parse_btn:
                _parse_all_panel_availability()

            st.markdown("#### Available Time Slots")

            intersections = st.session_state.get("computed_intersections", [])
            panel_interviewers = st.session_state.get("panel_interviewers", [])
            interviewer_count = len([i for i in panel_interviewers if i.get("slots")])

            if st.session_state["slots"]:
                # Filter mode selector (only show if multiple interviewers)
                if interviewer_count > 1:
                    from slot_intersection import filter_slots_by_availability

                    filter_col1, filter_col2 = st.columns([2, 1])
                    with filter_col1:
                        filter_options = [
                            ("all_available", f"All {interviewer_count} must be available"),
                            ("any_n", "At least N are available"),
                            ("show_all", "Show all slots"),
                        ]
                        filter_mode = st.selectbox(
                            "Show slots where:",
                            options=filter_options,
                            format_func=lambda x: x[1],
                            key="slot_filter_mode_select"
                        )
                        st.session_state["slot_filter_mode"] = filter_mode[0]

                    with filter_col2:
                        if filter_mode[0] == "any_n":
                            min_n = st.number_input(
                                "Minimum N",
                                min_value=1,
                                max_value=interviewer_count,
                                value=max(1, interviewer_count - 1),
                                key="slot_filter_min_n_input"
                            )
                            st.session_state["slot_filter_min_n"] = min_n

                    # Apply filter
                    filtered_slots = filter_slots_by_availability(
                        intersections,
                        st.session_state.get("slot_filter_mode", "all_available"),
                        st.session_state.get("slot_filter_min_n", 1),
                        interviewer_count
                    )
                else:
                    filtered_slots = st.session_state["slots"]

                if not filtered_slots:
                    st.warning("No slots match the current filter. Try relaxing the availability requirement.")
                    selected_slot = None
                else:
                    st.info("Select a slot to create an invite, or generate a candidate email.")

                    # Build slot labels with availability info
                    from slot_intersection import format_slot_label_with_availability

                    def get_slot_label(slot):
                        if interviewer_count > 1:
                            return format_slot_label_with_availability(slot, interviewer_count)
                        return format_slot_label(slot)

                    slot_labels = [get_slot_label(s) for s in filtered_slots]
                    selected_label = st.selectbox("Select slot", options=slot_labels, key="selected_slot_label")
                    selected_slot = filtered_slots[slot_labels.index(selected_label)]

                    # Show availability indicator for panel interviews
                    if interviewer_count > 1 and selected_slot:
                        avail = selected_slot.get("available_count", interviewer_count)
                        total = selected_slot.get("total_interviewers", interviewer_count)
                        available_names = selected_slot.get("available_names", [])

                        if avail == total:
                            st.success(f"All {total} interviewers available")
                        elif avail >= total * 0.75:
                            missing = [
                                i.get("name") or i.get("email")
                                for i in panel_interviewers
                                if i["id"] not in selected_slot.get("available_interviewers", [])
                                and i.get("slots")
                            ]
                            st.info(f"{avail}/{total} available. Missing: {', '.join(missing) if missing else 'None'}")
                        else:
                            st.warning(f"Only {avail}/{total} interviewers available: {', '.join(available_names)}")

                    # Real-time timezone conversion preview
                    if selected_slot:
                        from timezone_utils import safe_zoneinfo, from_utc, format_time_for_display
                        try:
                            # Parse the slot time as display timezone
                            slot_dt_naive = datetime.strptime(
                                f"{selected_slot['date']}T{selected_slot['start']}:00",
                                "%Y-%m-%dT%H:%M:%S"
                            )
                            zi, _ = safe_zoneinfo(tz_name, fallback="UTC")
                            slot_dt_local = slot_dt_naive.replace(tzinfo=zi)

                            # Convert to UTC for reference
                            from timezone_utils import to_utc
                            slot_utc = to_utc(slot_dt_local)

                            # Show conversion to common timezones
                            st.markdown("**Time Conversion Preview:**")
                            preview_tzs = ["UTC", "America/New_York", "America/Los_Angeles", "Europe/London", "Asia/Tokyo"]
                            # Add display timezone if not in list
                            if tz_name not in preview_tzs:
                                preview_tzs.insert(0, tz_name)

                            conversion_items = []
                            for preview_tz in preview_tzs:
                                try:
                                    converted = from_utc(slot_utc, preview_tz)
                                    time_str = converted.strftime("%a %b %d, %I:%M %p %Z")
                                    # Highlight the display timezone
                                    if preview_tz == tz_name:
                                        conversion_items.append(f"**{preview_tz}**: {time_str} *(selected)*")
                                    else:
                                        conversion_items.append(f"{preview_tz}: {time_str}")
                                except Exception:
                                    pass

                            st.caption(" | ".join(conversion_items[:4]))  # Show top 4
                        except (ValueError, TypeError):
                            pass  # Skip preview on invalid date

                    # DST Warning Check
                    if selected_slot:
                        from timezone_utils import is_near_dst_transition
                        try:
                            slot_date = datetime.strptime(selected_slot["date"], "%Y-%m-%d").date()
                            slot_dt = datetime.combine(slot_date, datetime.min.time())

                            # Check display timezone for DST transition
                            is_near, trans_date, trans_type = is_near_dst_transition(slot_dt, tz_name, days_threshold=7)
                            if is_near and trans_date:
                                direction = "spring forward" if trans_type == "spring_forward" else "fall back"
                                st.warning(
                                    f"DST Alert: Clocks {direction} on {trans_date.strftime('%B %d, %Y')} "
                                    f"in {tz_name}. Please verify the scheduled time."
                                )
                        except (ValueError, TypeError):
                            pass  # Skip DST check on invalid date
            else:
                st.info("No slots extracted yet. Upload availability and click Parse All Availability.")
                selected_slot = None

        with col_right:
            st.markdown("#### Candidate")
            candidate_name = st.text_input("Candidate Name", key="cand_name")
            candidate_email = st.text_input("Candidate Email (required)", key="cand_email")

            # Candidate timezone - pre-populate with inferred timezone from calendar
            inferred_tz = extract_common_timezone(st.session_state.get("slots", []))
            if inferred_tz:
                # Update session state if inference found a timezone
                st.session_state["candidate_timezone"] = inferred_tz

            candidate_tz_default = st.session_state.get("candidate_timezone", get_default_timezone())
            candidate_tz_idx = _common_timezones().index(candidate_tz_default) if candidate_tz_default in _common_timezones() else 0

            candidate_timezone = st.selectbox(
                "Candidate Timezone",
                options=_common_timezones(),
                index=candidate_tz_idx,
                key="candidate_timezone_select",
                help="Times in the invitation will be shown in this timezone"
            )

            if inferred_tz and inferred_tz == candidate_timezone:
                st.caption("Auto-detected from calendar screenshot")

            # Show candidate's view of the selected time
            if selected_slot:
                from timezone_utils import safe_zoneinfo, to_utc, from_utc, format_datetime_for_display
                try:
                    slot_dt_naive = datetime.strptime(
                        f"{selected_slot['date']}T{selected_slot['start']}:00",
                        "%Y-%m-%dT%H:%M:%S"
                    )
                    zi, _ = safe_zoneinfo(tz_name, fallback="UTC")
                    slot_dt_local = slot_dt_naive.replace(tzinfo=zi)
                    slot_utc = to_utc(slot_dt_local)

                    # Show what candidate will see
                    candidate_view = format_datetime_for_display(slot_utc, candidate_timezone)
                    st.success(f"Candidate will see: **{candidate_view}**")
                except (ValueError, TypeError):
                    pass

            st.markdown("#### Invite details")
            is_teams = st.selectbox("Interview type", options=["Teams", "Non-Teams"], index=0, key="interview_type") == "Teams"
            subject = st.text_input("Subject/title", value=f"Interview: {role_title}" if role_title else "Interview", key="subject")
            agenda = st.text_area("Description/agenda", value="Interview discussion.", key="agenda")
            location = st.text_input("Location (non-Teams)", value="", key="location")

            include_recruiter = st.checkbox("Include recruiter as attendee", value=False, key="include_recruiter")

            st.markdown("----")
            st.markdown("#### Actions")

            # Generate email to candidate (existing behavior)
            if st.button("Generate Candidate Scheduling Email"):
                body = build_scheduling_email(role_title, recruiter_name or "Recruiter", st.session_state["slots"])
                st.session_state["candidate_email_body"] = body

            if "candidate_email_body" in st.session_state and st.session_state["candidate_email_body"]:
                st.text_area("Email preview", st.session_state["candidate_email_body"], height=200)
                if st.button("Send Email"):
                    ok = send_email_graph(
                        subject=f"Interview availability: {role_title}",
                        body=st.session_state["candidate_email_body"],
                        to_emails=[candidate_email] if candidate_email else [],
                        cc_emails=[recruiter_email] if recruiter_email else None,
                    )
                    audit.log(
                        "graph_sent_scheduling_email" if ok else "graph_send_failed",
                        actor=recruiter_email or "",
                        candidate_email=candidate_email or "",
                        hiring_manager_email=hiring_manager_email or "",
                        recruiter_email=recruiter_email or "",
                        role_title=role_title or "",
                        payload={"subject": f"Interview availability: {role_title}"},
                        status="success" if ok else "failed",
                        error_message="" if ok else "Graph email send failed",
                    )
                    if ok:
                        st.success("Email sent.")
                    else:
                        st.error("Email send failed (see message above).")

            # Create Graph event
            # Collect panel interviewers from session state
            panel_interviewers_for_invite = [
                {"name": i.get("name", ""), "email": i.get("email", "")}
                for i in st.session_state.get("panel_interviewers", [])
                if i.get("email")  # Only include interviewers with valid emails
            ]

            # Determine if we have enough info to create invite
            has_interviewers = bool(panel_interviewers_for_invite) or bool(hiring_manager_email)
            create_disabled = not (selected_slot and has_interviewers and candidate_email)
            if st.button("Create & Send Interview Invite", disabled=create_disabled):
                _handle_create_invite(
                    audit=audit,
                    selected_slot=selected_slot,
                    tz_name=tz_name,
                    candidate_timezone=candidate_timezone,
                    duration_minutes=int(st.session_state["duration_minutes"]),
                    role_title=role_title,
                    subject=subject,
                    agenda=agenda,
                    location=location,
                    is_teams=is_teams,
                    candidate=(candidate_email, candidate_name),
                    hiring_manager=(hiring_manager_email, hiring_manager_name),
                    recruiter=(recruiter_email, recruiter_name),
                    include_recruiter=include_recruiter,
                    panel_interviewers=panel_interviewers_for_invite if panel_interviewers_for_invite else None,
                )

            # ICS fallback download button (available after generation)
            if st.session_state.get("last_invite_ics_bytes"):
                st.download_button(
                    "Download .ics (Add to calendar)",
                    data=st.session_state["last_invite_ics_bytes"],
                    file_name="powerdash_interview_invite.ics",
                    mime="text/calendar",
                )
                audit.log(
                    "ics_downloaded",
                    actor=recruiter_email or "",
                    candidate_email=candidate_email or "",
                    hiring_manager_email=hiring_manager_email or "",
                    recruiter_email=recruiter_email or "",
                    role_title=role_title or "",
                    event_id=st.session_state.get("last_graph_event_id", ""),
                    payload={"uid": st.session_state.get("last_invite_uid", "")},
                    status="success",
                )

                # Optional email ICS via Graph
                if st.button("Email .ics (optional)"):
                    ok = send_email_graph(
                        subject=subject,
                        body=agenda,
                        to_emails=[candidate_email, hiring_manager_email] + ([recruiter_email] if include_recruiter and recruiter_email else []),
                        attachment={
                            "data": st.session_state["last_invite_ics_bytes"],
                            "maintype": "text",
                            "subtype": "calendar",
                            "filename": "invite.ics",
                        },
                    )
                    audit.log(
                        "graph_sent_ics" if ok else "graph_send_failed",
                        actor=recruiter_email or "",
                        candidate_email=candidate_email or "",
                        hiring_manager_email=hiring_manager_email or "",
                        recruiter_email=recruiter_email or "",
                        role_title=role_title or "",
                        event_id=st.session_state.get("last_graph_event_id", ""),
                        payload={"uid": st.session_state.get("last_invite_uid", "")},
                        status="success" if ok else "failed",
                        error_message="" if ok else "Graph email send failed",
                    )
                    st.success("ICS emailed.") if ok else st.error("Failed to email ICS.")

    # ========= TAB: Scheduler Inbox =========
    with tab_inbox:
        st.subheader("Scheduler Inbox")
        st.caption("Reads unread emails from the scheduler mailbox via Microsoft Graph API.")
        emails, graph_error, is_configured = fetch_unread_emails_graph()
        if not is_configured:
            st.warning("Graph API is not configured. Add graph_tenant_id, graph_client_id, graph_client_secret, and graph_scheduler_mailbox to your secrets.")
        elif graph_error:
            st.error(f"Failed to fetch emails: {graph_error}")
        elif not emails:
            st.success("✓ Connected to mailbox. No unread emails found.")
        else:
            st.write(f"Found {len(emails)} unread email(s).")
            for i, e in enumerate(emails, start=1):
                with st.expander(f"{i}. {e['subject'] or '(no subject)'} — {e['from']}"):
                    st.write(e.get("date", ""))
                    body = e.get("body", "")
                    st.text_area("Body", body, height=160)

                    if st.session_state["slots"]:
                        choice = detect_slot_choice_from_text(body, st.session_state["slots"])
                        if choice:
                            st.success(f"Detected slot choice: {choice.get('date')} {choice.get('start')}")
                        else:
                            st.info("No slot choice detected from this email.")

    # ========= TAB: Calendar Invites =========
    with tab_invites:
        st.subheader("Calendar Invites")
        st.caption("Manage scheduled interviews (reschedule/cancel).")

        interviews = audit.list_interviews(limit=200)
        if not interviews:
            st.info("No interviews stored yet. Create an invite from the first tab.")
        else:
            # show compact table
            st.dataframe(
                [
                    {
                        "created_utc": r["created_utc"],
                        "role_title": r["role_title"],
                        "candidate": r["candidate_email"],
                        "hiring_manager": r["hiring_manager_email"],
                        "start_utc": r["start_utc"],
                        "graph_event_id": r["graph_event_id"],
                        "teams_join_url": (r["teams_join_url"][:45] + "…") if r.get("teams_join_url") else "",
                        "status": r.get("last_status", ""),
                    }
                    for r in interviews
                ],
                use_container_width=True,
                hide_index=True,
            )

            st.markdown("----")
            st.markdown("#### Reschedule / Cancel")

            event_ids = [r["graph_event_id"] for r in interviews if r.get("graph_event_id")]
            selected_event_id = st.selectbox("Select event", options=event_ids)
            selected_row = next((r for r in interviews if r.get("graph_event_id") == selected_event_id), None)

            if selected_row:
                display_tz = st.selectbox("Display timezone", options=_common_timezones(), index=_tz_index(selected_row.get("display_timezone")))
                st.write(f"Current start (UTC): {selected_row.get('start_utc')}")
                st.write(f"Current end (UTC): {selected_row.get('end_utc')}")
                try:
                    start_local = from_utc(datetime.fromisoformat(selected_row["start_utc"]), display_tz)
                except Exception:
                    start_local = None

                new_date = st.date_input("New date", value=start_local.date() if start_local else datetime.now().date())
                new_time = st.time_input("New time", value=start_local.time().replace(second=0, microsecond=0) if start_local else datetime.now().time().replace(second=0, microsecond=0))
                new_duration = st.number_input("Duration (minutes)", min_value=15, max_value=240, step=15, value=int(selected_row.get("duration_minutes") or 30))

                colA, colB = st.columns(2)
                with colA:
                    if st.button("Reschedule", type="primary"):
                        _handle_reschedule(
                            audit=audit,
                            event_id=selected_event_id,
                            new_date=new_date,
                            new_time=new_time,
                            duration_minutes=int(new_duration),
                            tz_name=display_tz,
                            context_row=selected_row,
                        )
                with colB:
                    if st.button("Cancel", type="secondary"):
                        _handle_cancel(audit=audit, event_id=selected_event_id, context_row=selected_row)

    # ========= TAB: Audit Log =========
    with tab_audit:
        st.subheader("Audit Log")
        st.caption("Append-only log of scheduling actions.")
        rows = audit.list_recent_audit(limit=300)
        if not rows:
            st.info("No audit entries yet.")
        else:
            st.dataframe(
                [
                    {
                        "timestamp_utc": r["timestamp_utc"],
                        "action": r["action"],
                        "status": r["status"],
                        "candidate": r["candidate_email"],
                        "hiring_manager": r["hiring_manager_email"],
                        "event_id": r["event_id"],
                        "error": (r["error_message"][:80] + "…") if r.get("error_message") and len(r["error_message"]) > 80 else (r.get("error_message") or ""),
                    }
                    for r in rows
                ],
                use_container_width=True,
                hide_index=True,
            )

            with st.expander("Show raw payload for a row"):
                idx = st.number_input("Row index (0 = most recent)", min_value=0, max_value=max(0, len(rows) - 1), value=0)
                st.json(rows[int(idx)])

    # ========= TAB: Graph Diagnostics =========
    with tab_diag:
        st.subheader("Graph Diagnostics")
        st.caption("Use this to verify Graph auth and mailbox access. No secrets are displayed.")

        cfg = get_graph_config()
        if not cfg:
            st.warning("Graph is not configured. Add graph_tenant_id, graph_client_id, graph_client_secret, graph_scheduler_mailbox in Streamlit secrets.")
        else:
            st.code(f"Scheduler mailbox: {cfg.scheduler_mailbox}")
            client = GraphClient(cfg)

            if st.button("Test token acquisition"):
                try:
                    token = client.get_token(force_refresh=True)
                    st.success(f"Token OK (length={len(token)})")
                    audit.log("graph_token_ok", payload={"length": len(token)}, status="success")
                except Exception as e:
                    st.error(str(e))
                    audit.log("graph_token_failed", payload={"error": str(e)}, status="failed", error_message=str(e))

            if st.button("Test calendar read (top 5)"):
                try:
                    data = client.test_calendar_read(top=5)
                    st.success("Calendar read OK.")
                    st.json(data)
                    audit.log("graph_calendar_read_ok", payload={"top": 5}, status="success")
                except GraphAPIError as e:
                    st.error(f"{e} — details below")
                    st.json(e.response_json)
                    audit.log("graph_calendar_read_failed", payload=e.response_json, status="failed", error_message=str(e))

            st.markdown("----")
            st.markdown("#### Dummy event (optional)")
            dry_run = st.checkbox("Dry run (do not create)", value=True)
            tz_name = st.selectbox("Timezone", options=_common_timezones(), index=_common_timezones().index(get_default_timezone()) if get_default_timezone() in _common_timezones() else 0)
            dt = datetime.now().replace(second=0, microsecond=0) + timedelta(hours=2)
            date = st.date_input("Date", value=dt.date())
            time = st.time_input("Time", value=dt.time())
            start_local = datetime.combine(date, time).replace(tzinfo=_zoneinfo(tz_name))
            end_local = start_local + timedelta(minutes=30)

            if st.button("Create dummy event"):
                try:
                    start_dt = {"dateTime": start_local.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": tz_name}
                    end_dt = {"dateTime": end_local.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": tz_name}
                    out = client.create_dummy_event("PowerDash Diagnostics", start_dt, end_dt, dry_run=dry_run)
                    st.success("OK")
                    st.json(out)
                    audit.log("graph_dummy_event_ok", payload={"dry_run": dry_run}, status="success")
                except GraphAPIError as e:
                    st.error(f"{e}")
                    st.json(e.response_json)
                    audit.log("graph_dummy_event_failed", payload=e.response_json, status="failed", error_message=str(e))


# ----------------------------
# Internal UI handlers
# ----------------------------
def _parse_availability_upload(upload) -> List[Dict[str, str]]:
    data = upload.read()
    name = (upload.name or "").lower()
    slots: List[Dict[str, str]] = []

    if name.endswith(".pdf"):
        imgs = pdf_to_images(data, max_pages=3)
        for img in imgs:
            slots.extend(parse_slots_from_image(img))

    elif name.endswith(".docx"):
        # Strategy: Extract text + tables, then also check embedded images

        # 1. Parse text content (paragraphs + tables)
        text = docx_to_text(data)
        if text:
            slots.extend(parse_slots_from_text(text))

        # 2. Extract and parse embedded images (optional enhancement)
        embedded_images = docx_extract_images(data, max_images=3)
        for img in embedded_images:
            slots.extend(parse_slots_from_image(img))

    else:
        # Assume image file (png, jpg, jpeg)
        img = Image.open(io.BytesIO(data)).convert("RGB")
        slots.extend(parse_slots_from_image(img))

    # De-duplicate slots by (date, start, end) tuple
    uniq = {(s["date"], s["start"], s["end"]): s for s in slots}
    return list(uniq.values())


def _parse_all_panel_availability() -> None:
    """Parse availability for all interviewers and compute intersection.

    Handles both uploaded files and manually-entered slots.
    """
    from slot_intersection import (
        normalize_slots_to_utc,
        merge_adjacent_slots,
        compute_intersection,
    )

    interviewers = st.session_state.get("panel_interviewers", [])
    tz_name = st.session_state["selected_timezone"]
    min_duration = st.session_state["duration_minutes"]

    all_interviewer_slots: Dict[int, List] = {}
    interviewer_names: Dict[int, str] = {}
    parse_errors = []
    total_uploaded = 0
    total_manual = 0

    for interviewer in interviewers:
        # Get existing manual slots (preserve them)
        existing_manual_slots = [s for s in interviewer.get("slots", []) if s.get("source") == "manual"]
        total_manual += len(existing_manual_slots)

        try:
            if interviewer.get("file"):
                # Reset file position before reading
                interviewer["file"].seek(0)
                # Parse the uploaded file
                uploaded_slots = _parse_availability_upload(interviewer["file"])
                # Mark uploaded slots with source
                for s in uploaded_slots:
                    s["source"] = "uploaded"
                total_uploaded += len(uploaded_slots)
                # Merge manual + uploaded, preferring manual for duplicates
                interviewer["slots"] = _merge_slots(existing_manual_slots, uploaded_slots)
            elif existing_manual_slots:
                # No file but has manual slots - keep them
                interviewer["slots"] = existing_manual_slots

            # Include interviewer if they have any slots
            if interviewer.get("slots"):
                # Build interviewer name for display
                name = interviewer.get("name") or interviewer.get("email") or f"Interviewer {interviewer['id']}"
                interviewer_names[interviewer["id"]] = name

                # Normalize to UTC for intersection
                normalized = normalize_slots_to_utc(interviewer["slots"], tz_name)
                merged = merge_adjacent_slots(normalized)
                all_interviewer_slots[interviewer["id"]] = merged

        except Exception as e:
            interviewer_name = interviewer.get("name") or f"Interviewer {interviewer.get('id', '?')}"
            parse_errors.append(f"{interviewer_name}: {e}")

    if parse_errors:
        for err in parse_errors:
            st.error(err)

    # Compute intersection
    if all_interviewer_slots:
        intersections = compute_intersection(
            all_interviewer_slots,
            min_duration_minutes=min_duration,
            display_timezone=tz_name,
            interviewer_names=interviewer_names,
        )
        st.session_state["computed_intersections"] = intersections

        # Also update legacy "slots" for backward compatibility
        st.session_state["slots"] = intersections

        num_interviewers = len(all_interviewer_slots)
        total_slots = total_uploaded + total_manual

        if num_interviewers == 1:
            if total_manual > 0 and total_uploaded > 0:
                st.success(f"Processed {total_slots} slot(s) ({total_manual} manual, {total_uploaded} uploaded).")
            elif total_manual > 0:
                st.success(f"Processed {total_manual} manual slot(s).")
            else:
                st.success(f"Extracted {total_uploaded} slot(s) from uploaded file.")
        else:
            full_overlap = sum(1 for s in intersections if s.get("is_full_overlap", False))
            source_info = []
            if total_manual > 0:
                source_info.append(f"{total_manual} manual")
            if total_uploaded > 0:
                source_info.append(f"{total_uploaded} uploaded")
            source_str = f" ({', '.join(source_info)})" if source_info else ""
            st.success(
                f"Processed {total_slots} total slots{source_str} from {num_interviewers} interviewers. "
                f"Found {len(intersections)} intersection slot(s) ({full_overlap} with all available)."
            )
    else:
        st.warning("No availability found. Please upload calendars or add slots manually.")


def _zoneinfo(tz_name: str):
    """Get ZoneInfo with validation. Falls back to UTC if invalid."""
    zi, was_valid = safe_zoneinfo(tz_name, fallback="UTC")
    if not was_valid:
        st.warning(f"Invalid timezone '{tz_name}', using UTC")
    return zi


def _common_timezones() -> List[str]:
    # Keep concise; you can expand later.
    return [
        "UTC",
        "Europe/London",
        "Europe/Dublin",
        "Europe/Paris",
        "Europe/Rome",
        "Europe/Berlin",
        "America/New_York",
        "America/Chicago",
        "America/Denver",
        "America/Los_Angeles",
        "America/Toronto",
        "America/Sao_Paulo",
        "Asia/Dubai",
        "Asia/Kolkata",
        "Asia/Singapore",
        "Asia/Tokyo",
        "Australia/Sydney",
    ]


def _tz_index(tz_name: str | None) -> int:
    tzs = _common_timezones()
    if tz_name and tz_name in tzs:
        return tzs.index(tz_name)
    return tzs.index(get_default_timezone()) if get_default_timezone() in tzs else 0


def _handle_create_invite(
    *,
    audit: AuditLog,
    selected_slot: Dict[str, str],
    tz_name: str,
    candidate_timezone: str,
    duration_minutes: int,
    role_title: str,
    subject: str,
    agenda: str,
    location: str,
    is_teams: bool,
    candidate: Tuple[str, str],
    hiring_manager: Tuple[str, str],
    recruiter: Tuple[str, str],
    include_recruiter: bool,
    panel_interviewers: Optional[List[Dict[str, str]]] = None,
) -> None:
    candidate_email_raw, candidate_name = candidate
    hm_email_raw, hm_name = hiring_manager
    rec_email_raw, rec_name = recruiter

    # === INPUT VALIDATION ===
    # Validate timezones
    if not is_valid_timezone(tz_name):
        st.warning(f"Invalid display timezone '{tz_name}', using UTC")
        tz_name = "UTC"

    if not is_valid_timezone(candidate_timezone):
        st.warning(f"Invalid candidate timezone '{candidate_timezone}', using display timezone")
        candidate_timezone = tz_name

    # Validate emails
    try:
        candidate_email = validate_email(candidate_email_raw, "Candidate email")
        hm_email = validate_email(hm_email_raw, "Hiring manager email")
        rec_email = validate_email_optional(rec_email_raw, "Recruiter email")
    except ValidationError as e:
        st.error(f"Validation error: {e.message}")
        return

    # Validate slot format
    try:
        validate_slot(selected_slot)
    except ValidationError as e:
        st.error(f"Invalid time slot: {e.message}")
        return

    # Parse selected slot into a local datetime
    try:
        start_local_naive = datetime.fromisoformat(f"{selected_slot['date']}T{selected_slot['start']}:00")
    except ValueError as e:
        st.error(f"Selected slot has invalid date/time format: {e}")
        return

    zi, _ = safe_zoneinfo(tz_name, fallback="UTC")
    start_local = start_local_naive.replace(tzinfo=zi)
    end_local = start_local + timedelta(minutes=duration_minutes)

    start_utc = to_utc(start_local)
    end_utc = to_utc(end_local)

    # === IDEMPOTENCY CHECK ===
    existing = audit.interview_exists(
        candidate_email=candidate_email,
        hiring_manager_email=hm_email,
        role_title=role_title,
        start_utc=iso_utc(start_utc),
    )
    if existing:
        st.warning(
            f"An interview already exists for this candidate at this time. "
            f"Event ID: {existing.get('graph_event_id', 'N/A')}"
        )
        # Use a unique key based on the slot to avoid Streamlit duplicate key errors
        checkbox_key = f"force_dup_{selected_slot['date']}_{selected_slot['start']}"
        if not st.checkbox("Create duplicate anyway?", key=checkbox_key):
            return

    attendees: List[Tuple[str, str]] = [(candidate_email, candidate_name)]

    # Build attendees from panel interviewers if provided, otherwise use hiring manager
    is_panel = panel_interviewers and len(panel_interviewers) > 1
    validated_panel: List[Dict[str, str]] = []

    if panel_interviewers:
        seen_emails = {candidate_email}  # Avoid duplicating candidate
        for pi in panel_interviewers:
            pi_email = (pi.get("email") or "").strip().lower()
            if pi_email and pi_email not in seen_emails:
                try:
                    validated_email = validate_email(pi_email, "Panel interviewer email")
                    validated_panel.append({"name": pi.get("name", ""), "email": validated_email})
                    attendees.append((validated_email, pi.get("name", "")))
                    seen_emails.add(validated_email)
                except ValidationError:
                    pass  # Skip invalid emails
    else:
        # Fall back to single hiring manager (backward compatibility)
        attendees.append((hm_email, hm_name))

    if include_recruiter and rec_email:
        attendees.append((rec_email, rec_name))

    organizer_email = str(get_secret("graph_scheduler_mailbox", "scheduling@powerdashhr.com"))
    organizer_name = "PowerDash Scheduler"

    # Update subject for panel interviews
    effective_subject = subject
    if is_panel:
        if not subject.startswith("Panel Interview"):
            effective_subject = f"Panel Interview: {role_title} - {candidate_name}"

    # Always generate ICS (so we have a fallback even if Graph works)
    ics_bytes = _build_ics(
        organizer_email=organizer_email,
        organizer_name=organizer_name,
        attendee_emails=[a[0] for a in attendees],
        summary=effective_subject,
        description=agenda,
        dtstart_utc=start_utc,
        dtend_utc=end_utc,
        location=("Microsoft Teams" if is_teams else (location or "Interview")),
        url="",
        uid_hint=f"{role_title}|{candidate_email}|{hm_email}",
        display_timezone=candidate_timezone,
    )
    st.session_state["last_invite_ics_bytes"] = ics_bytes
    st.session_state["last_invite_uid"] = stable_uid(f"{role_title}|{candidate_email}|{hm_email}", organizer_email, start_utc.isoformat())
    audit.log(
        "ics_generated",
        actor=rec_email or "",
        candidate_email=candidate_email,
        hiring_manager_email=hm_email,
        recruiter_email=rec_email or "",
        role_title=role_title,
        payload={"uid": st.session_state["last_invite_uid"]},
        status="success",
    )

    client = _make_graph_client()
    if not client:
        st.warning("Graph is not configured. Using .ics fallback only.")
        return

    # Format time display for candidate's timezone
    from timezone_utils import format_datetime_for_display
    candidate_time_display = format_datetime_for_display(start_utc, candidate_timezone)

    # Build body with candidate-friendly time display
    body_html = f"<p><strong>Interview Time (your timezone): {candidate_time_display}</strong></p>"

    # Add panel members to body if panel interview
    if is_panel and validated_panel:
        body_html += "<p><strong>Interview Panel:</strong></p><ul>"
        for pi in validated_panel:
            name = pi.get("name") or pi.get("email", "")
            body_html += f"<li>{name}</li>"
        body_html += "</ul>"

    if agenda:
        body_html += f"<p>{agenda.replace(chr(10), '<br>')}</p>"

    payload = _graph_event_payload(
        subject=effective_subject,
        body_html=body_html,
        start_local=start_local,
        end_local=end_local,
        time_zone=candidate_timezone,  # Use candidate timezone for calendar event
        attendees=attendees,
        is_teams=is_teams,
        location=location,
    )

    try:
        created = client.create_event(payload)
        event_id = created.get("id", "")
        teams_url = ""
        if is_teams:
            teams_url = (created.get("onlineMeeting") or {}).get("joinUrl") or ""
        st.session_state["last_graph_event_id"] = event_id
        st.session_state["last_teams_join_url"] = teams_url

        # Re-generate ICS including Teams URL if present (better fallback)
        if teams_url:
            st.session_state["last_invite_ics_bytes"] = _build_ics(
                organizer_email=organizer_email,
                organizer_name=organizer_name,
                attendee_emails=[a[0] for a in attendees],
                summary=effective_subject,
                description=agenda,
                dtstart_utc=start_utc,
                dtend_utc=end_utc,
                location="Microsoft Teams",
                url=teams_url,
                uid_hint=f"{role_title}|{candidate_email}|{hm_email}",
                display_timezone=candidate_timezone,
            )

        audit.log(
            "graph_create_event",
            actor=rec_email or "",
            candidate_email=candidate_email,
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            role_title=role_title,
            event_id=event_id,
            payload=payload,
            status="success",
        )

        # Serialize panel interviewers for database storage
        panel_json = ""
        if validated_panel:
            import json as _json
            panel_json = _json.dumps(validated_panel)

        audit.upsert_interview(
            role_title=role_title,
            candidate_email=candidate_email,
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            duration_minutes=duration_minutes,
            start_utc=iso_utc(start_utc),
            end_utc=iso_utc(end_utc),
            display_timezone=tz_name,
            candidate_timezone=candidate_timezone,
            graph_event_id=event_id,
            teams_join_url=teams_url,
            subject=effective_subject,
            last_status="created",
            panel_interviewers_json=panel_json,
            is_panel_interview=is_panel,
        )

        st.success("Invite created and sent via Microsoft Graph.")
        if teams_url:
            st.link_button("Open Teams meeting link", teams_url)
    except (GraphAuthError, GraphAPIError) as e:
        details = getattr(e, "response_json", None)
        st.error("Graph scheduling failed. .ics fallback is available for download.")
        if details:
            st.json(details)
        audit.log(
            "graph_create_failed",
            actor=rec_email or "",
            candidate_email=candidate_email,
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            role_title=role_title,
            payload={"error": str(e), "details": details},
            status="failed",
            error_message=str(e),
        )


def _handle_reschedule(
    *,
    audit: AuditLog,
    event_id: str,
    new_date,
    new_time,
    duration_minutes: int,
    tz_name: str,
    context_row: Dict[str, Any],
) -> None:
    client = _make_graph_client()
    if not client:
        st.error("Graph is not configured.")
        return

    start_local = datetime.combine(new_date, new_time).replace(tzinfo=_zoneinfo(tz_name))
    end_local = start_local + timedelta(minutes=duration_minutes)
    start_utc = to_utc(start_local)
    end_utc = to_utc(end_local)

    patch = {
        "start": {"dateTime": start_local.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": tz_name},
        "end": {"dateTime": end_local.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": tz_name},
    }

    try:
        client.patch_event(event_id, patch, send_updates="all")
        audit.log(
            "graph_reschedule_event",
            actor=context_row.get("recruiter_email", "") or "",
            candidate_email=context_row.get("candidate_email", "") or "",
            hiring_manager_email=context_row.get("hiring_manager_email", "") or "",
            recruiter_email=context_row.get("recruiter_email", "") or "",
            role_title=context_row.get("role_title", "") or "",
            event_id=event_id,
            payload=patch,
            status="success",
        )
        st.success("Event rescheduled. Attendees should receive updated invites.")
    except GraphAPIError as e:
        st.error("Reschedule failed.")
        st.json(e.response_json)
        audit.log(
            "graph_reschedule_failed",
            actor=context_row.get("recruiter_email", "") or "",
            candidate_email=context_row.get("candidate_email", "") or "",
            hiring_manager_email=context_row.get("hiring_manager_email", "") or "",
            recruiter_email=context_row.get("recruiter_email", "") or "",
            role_title=context_row.get("role_title", "") or "",
            event_id=event_id,
            payload=e.response_json,
            status="failed",
            error_message=str(e),
        )


def _handle_cancel(*, audit: AuditLog, event_id: str, context_row: Dict[str, Any]) -> None:
    client = _make_graph_client()
    if not client:
        st.error("Graph is not configured.")
        return

    try:
        client.delete_event(event_id)
        audit.log(
            "graph_cancel_event",
            actor=context_row.get("recruiter_email", "") or "",
            candidate_email=context_row.get("candidate_email", "") or "",
            hiring_manager_email=context_row.get("hiring_manager_email", "") or "",
            recruiter_email=context_row.get("recruiter_email", "") or "",
            role_title=context_row.get("role_title", "") or "",
            event_id=event_id,
            payload={"event_id": event_id},
            status="success",
        )
        st.success("Event cancelled. Attendees should receive cancellation notices.")
    except GraphAPIError as e:
        st.error("Cancel failed.")
        st.json(e.response_json)
        audit.log(
            "graph_cancel_failed",
            actor=context_row.get("recruiter_email", "") or "",
            candidate_email=context_row.get("candidate_email", "") or "",
            hiring_manager_email=context_row.get("hiring_manager_email", "") or "",
            recruiter_email=context_row.get("recruiter_email", "") or "",
            role_title=context_row.get("role_title", "") or "",
            event_id=event_id,
            payload=e.response_json,
            status="failed",
            error_message=str(e),
        )


if __name__ == "__main__":
    main()
